from __future__ import annotations

import logging
import shutil
import subprocess
import time
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt, RGBColor

from cryodaq.report_process import terminate_descendant_tree
from cryodaq.report_state import (
    ReportContractError,
    compute_source_fingerprint,
    resolve_experiment_dir,
    resolve_report_paths,
)
from cryodaq.reporting.data import ReportDataExtractor
from cryodaq.reporting.sections import SECTION_REGISTRY

# Hard wall-clock bound for the best-effort LibreOffice PDF conversion.
# A hung ``soffice`` (stale lock, headless-profile contention, second
# instance) would otherwise block the worker thread forever — and a Python
# thread cannot be killed mid-call — eventually exhausting the thread pool.
# On timeout the report degrades to docx-only, exactly like a missing soffice.
_SOFFICE_TIMEOUT_S = 120

logger = logging.getLogger(__name__)


@dataclass(frozen=True, slots=True)
class ReportGenerationResult:
    docx_path: Path
    pdf_path: Path | None
    assets_dir: Path
    sections: tuple[str, ...]
    skipped: bool = False
    reason: str = ""


class ReportGenerator:
    _BASE_RAW_SECTIONS = (
        "title_page",
        "experiment_metadata_section",
        "run_timeline_section",
        "run_parameters_section",
        "result_tables_section",
        "conductivity_section",
        "cooldown_section",
        "thermal_section",
        "pressure_section",
        "artifact_manifest_section",
    )
    _EDITABLE_ONLY_SECTIONS = (
        "operator_comments_section",
        "operator_interpretation_section",
        "operator_photos_section",
    )

    def __init__(self, data_dir: Path) -> None:
        self._data_dir = data_dir
        self._experiments_dir = data_dir / "experiments"
        self._extractor = ReportDataExtractor(data_dir)

    def generate(self, experiment_id: str) -> ReportGenerationResult:
        experiment_root = resolve_experiment_dir(self._data_dir, experiment_id)
        compute_source_fingerprint(experiment_root)
        report_paths = resolve_report_paths(experiment_root, create_reports=True)
        return self._generate_into(experiment_root, report_paths.reports)

    def generate_to_directory(
        self,
        experiment_id: str,
        output_dir: Path,
        *,
        deadline_epoch: float | None = None,
    ) -> ReportGenerationResult:
        """Render into one validated, child-owned generation staging directory."""
        experiment_root = resolve_experiment_dir(self._data_dir, experiment_id)
        report_paths = resolve_report_paths(experiment_root)
        staging_root = report_paths.staging_root
        output_dir = Path(output_dir)
        if output_dir.is_symlink():
            raise ReportContractError("report output directory must not be a symlink")
        resolved_output = output_dir.resolve()
        if (
            resolved_output.parent != staging_root
            or report_paths.experiment_root not in resolved_output.parents
        ):
            raise ReportContractError(
                "report output must be one generation staging directory"
            )
        return self._generate_into(
            experiment_root,
            resolved_output,
            deadline_epoch=deadline_epoch,
        )

    def _generate_into(
        self,
        experiment_root: Path,
        reports_dir: Path,
        *,
        deadline_epoch: float | None = None,
    ) -> ReportGenerationResult:
        metadata_path = experiment_root / "metadata.json"
        dataset = self._extractor.load_dataset(metadata_path)
        experiment = dataset.metadata["experiment"]
        template = dataset.metadata["template"]
        assets_dir = reports_dir / "assets"
        editable_docx_path = reports_dir / "report_editable.docx"
        raw_source_docx_path = reports_dir / "report_raw.docx"
        raw_pdf_path = reports_dir / "report_raw.pdf"

        if not bool(experiment.get("report_enabled", template.get("report_enabled", True))):
            return ReportGenerationResult(
                docx_path=editable_docx_path,
                pdf_path=None,
                assets_dir=assets_dir,
                sections=tuple(),
                skipped=True,
                reason="Формирование отчёта отключено шаблоном.",
            )

        reports_dir.mkdir(parents=True, exist_ok=True)
        assets_dir.mkdir(parents=True, exist_ok=True)

        raw_sections = self._resolve_raw_sections(dataset.metadata)
        editable_sections = tuple(list(raw_sections) + list(self._EDITABLE_ONLY_SECTIONS))

        # B1 (2026-07): Гемма-generated annotation used to be produced
        # in-process here (import + synchronous Ollama call from the
        # engine process). B1 moved all LLM/RAG code out of the engine —
        # the intro paragraph is not reinstated via a cross-process call
        # in this pass (see scratchpad/montana/exec/impl_b1.md, "forks").
        # ``_build_document`` already treats ``gemma_intro=None`` as
        # "skip the annotation section", which is exactly the existing
        # graceful-degradation behaviour when Ollama was unavailable.
        gemma_intro: str | None = None

        raw_document = self._build_document(dataset, assets_dir, raw_sections, gemma_intro)
        raw_document.save(str(raw_source_docx_path))
        pdf_path = self._try_convert_pdf(
            raw_source_docx_path,
            raw_pdf_path,
            deadline_epoch=deadline_epoch,
        )

        editable_document = self._build_document(
            dataset, assets_dir, editable_sections, gemma_intro
        )
        editable_document.save(str(editable_docx_path))

        return ReportGenerationResult(
            docx_path=editable_docx_path,
            pdf_path=pdf_path,
            assets_dir=assets_dir,
            sections=editable_sections,
        )

    def _build_document(
        self,
        dataset,
        assets_dir: Path,
        sections: tuple[str, ...],
        gemma_intro: str | None = None,
    ) -> Document:
        document = Document()
        self._apply_gost_formatting(document)

        from cryodaq.reporting.sections import _reset_counters

        _reset_counters()

        # Only break before major sections, not after every one
        _PAGE_BREAK_BEFORE = {
            "cooldown_section",
            "thermal_section",
            "pressure_section",
            "operator_log_section",
            "alarms_section",
            "config_section",
            "operator_comments_section",
            "artifact_manifest_section",
        }
        for index, section_name in enumerate(sections):
            if index > 0 and section_name in _PAGE_BREAK_BEFORE:
                document.add_page_break()
            renderer = SECTION_REGISTRY[section_name]
            renderer(document, dataset, assets_dir)
            # Insert Гемма annotation immediately after title page
            if section_name == "title_page" and gemma_intro:
                self._render_gemma_annotation(document, gemma_intro)
        return document

    @staticmethod
    def _render_gemma_annotation(document: Document, intro_text: str) -> None:
        """Insert Slice C Гемма-generated annotation section after the title page."""
        from cryodaq.utils.xml_safe import xml_safe

        document.add_heading("Аннотация", level=1)
        for para in intro_text.strip().split("\n"):
            if para.strip():
                document.add_paragraph(xml_safe(para.strip()))
        # Auto-generated marker as italicised note
        note = document.add_paragraph()
        run = note.add_run("Аннотация сгенерирована автоматически: Гемма (gemma4:e4b).")
        run.italic = True
        run.font.size = Pt(11)

    @staticmethod
    def _apply_gost_formatting(document: Document) -> None:
        """Apply ГОСТ Р 2.105-2019 page setup and styles."""
        # Page setup: A4, margins
        sec = document.sections[0]
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(30)
        sec.right_margin = Mm(15)
        sec.top_margin = Mm(20)
        sec.bottom_margin = Mm(20)

        # Normal style: Times New Roman 14pt, 1.5 spacing, first-line indent
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        pf = style.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.5
        pf.first_line_indent = Cm(1.25)

        # Heading 1: 16pt bold black
        h1 = document.styles["Heading 1"]
        h1.font.name = "Times New Roman"
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.first_line_indent = Cm(0)

        # Heading 2: 14pt bold black
        h2 = document.styles["Heading 2"]
        h2.font.name = "Times New Roman"
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(6)
        h2.paragraph_format.first_line_indent = Cm(0)

        # Title: 18pt bold centered
        ts = document.styles["Title"]
        ts.font.name = "Times New Roman"
        ts.font.size = Pt(18)
        ts.font.bold = True
        ts.font.color.rgb = RGBColor(0, 0, 0)
        ts.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ts.paragraph_format.first_line_indent = Cm(0)

        # Force black headings — clear theme color that overrides rgb
        from docx.oxml.ns import qn as _qn

        for sn in ("Heading 1", "Heading 2", "Title"):
            rpr = document.styles[sn].element.find(_qn("w:rPr"))
            if rpr is not None:
                ce = rpr.find(_qn("w:color"))
                if ce is not None:
                    ce.set(_qn("w:val"), "000000")
                    for attr in ("w:themeColor", "w:themeShade", "w:themeTint"):
                        if _qn(attr) in ce.attrib:
                            del ce.attrib[_qn(attr)]

        # List Bullet
        if "List Bullet" in document.styles:
            lb = document.styles["List Bullet"]
            lb.font.name = "Times New Roman"
            lb.font.size = Pt(14)
            lb.paragraph_format.line_spacing = 1.5

        # Footer with page number (center)
        footer = sec.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml.ns import qn

        run = fp.add_run()
        run._element.append(
            run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        )
        run2 = fp.add_run()
        instr = run2._element.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
        instr.text = " PAGE "
        run2._element.append(instr)
        run3 = fp.add_run()
        run3._element.append(
            run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        )

    def _resolve_raw_sections(self, metadata: dict) -> tuple[str, ...]:
        template = metadata.get("template", {})
        configured = [
            name for name in list(template.get("report_sections") or []) if name in SECTION_REGISTRY
        ]
        ordered: list[str] = []
        for name in self._BASE_RAW_SECTIONS + tuple(configured):
            if name not in ordered:
                ordered.append(name)
        if "config_section" not in ordered and "config_section" in SECTION_REGISTRY:
            ordered.append("config_section")
        return tuple(ordered)

    def _try_convert_pdf(
        self,
        source_docx_path: Path,
        target_pdf_path: Path,
        *,
        deadline_epoch: float | None = None,
    ) -> Path | None:
        """Best-effort DOCX→PDF via LibreOffice; None on any degradation.

        Every degradation path (missing soffice, timeout, failed conversion)
        is logged LOUD naming the consequence and the remedy. The absence is
        also surfaced to the operator: ``generate()`` returns ``pdf_path=None``,
        the engine report dict carries ``pdf_path: null``, and the archive
        panel renders "PDF: нет" (gui/shell/overlays/archive_panel.py).
        """
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            logger.warning(
                "ReportGenerator: soffice/LibreOffice не найден — PDF не создан, "
                "доступен только DOCX. Для PDF-конвертации установите LibreOffice."
            )
            return None
        output_dir = source_docx_path.parent
        timeout_s = float(_SOFFICE_TIMEOUT_S)
        if deadline_epoch is not None:
            # Reserve time for the editable document, hashing, promotion, and reply.
            timeout_s = min(timeout_s, max(0.1, deadline_epoch - time.time() - 2.0))
        command = [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            str(source_docx_path),
            "--outdir",
            str(output_dir),
        ]
        try:
            if deadline_epoch is None:
                # Preserve the direct synchronous API's characterized behavior.
                subprocess.run(
                    command,
                    check=False,
                    capture_output=True,
                    timeout=timeout_s,
                )
            else:
                process = subprocess.Popen(
                    command,
                    stdin=subprocess.DEVNULL,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    close_fds=True,
                )
                try:
                    process.wait(timeout=timeout_s)
                except subprocess.TimeoutExpired:
                    terminate_descendant_tree(process.pid)
                    process.wait(timeout=2.0)
                    raise
        except subprocess.TimeoutExpired:
            logger.error(
                "ReportGenerator: конвертация soffice в PDF превысила таймаут %d с — "
                "PDF не создан, доступен только DOCX. Проверьте зависшие процессы "
                "soffice / установку LibreOffice.",
                timeout_s,
            )
            return None
        produced = output_dir / f"{source_docx_path.stem}.pdf"
        if not produced.exists():
            logger.error(
                "ReportGenerator: soffice завершился, но PDF (%s) не создан — "
                "доступен только DOCX. Проверьте установку LibreOffice.",
                produced.name,
            )
            return None
        if produced != target_pdf_path:
            if target_pdf_path.exists():
                target_pdf_path.unlink()
            produced.replace(target_pdf_path)
        return target_pdf_path if target_pdf_path.exists() else None
