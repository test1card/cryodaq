"""Report section renderers for CryoDAQ experiment reports.

Each function has signature: (document: Document, dataset: ReportDataset, assets_dir: Path) -> None
"""

from __future__ import annotations

import csv
import math
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from pathlib import Path
from statistics import mean
from typing import Any, Callable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Mm, Pt

from cryodaq.reporting.data import HistoricalReading, ReportDataset

SectionRenderer = Callable[[Document, ReportDataset, Path], None]

# ---------------------------------------------------------------------------
# ГОСТ counters
# ---------------------------------------------------------------------------

_figure_counter = 0
_table_counter = 0


def _reset_counters() -> None:
    global _figure_counter, _table_counter
    _figure_counter = 0
    _table_counter = 0


def _next_figure() -> int:
    global _figure_counter
    _figure_counter += 1
    return _figure_counter


def _next_table() -> int:
    global _table_counter
    _table_counter += 1
    return _table_counter


# ---------------------------------------------------------------------------
# Global helpers
# ---------------------------------------------------------------------------

_MONTHS_RU = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}


def _format_dt(raw: Any, *, time_only: bool = False) -> str:
    """Format datetime as '22 марта 2026, 15:42' or '15:42:05'."""
    if raw is None:
        return "—"
    if isinstance(raw, str):
        raw = raw.strip()
        if not raw:
            return "—"
        try:
            if raw.endswith("Z"):
                raw = raw[:-1] + "+00:00"
            raw = datetime.fromisoformat(raw)
        except ValueError:
            return str(raw)
    if not isinstance(raw, datetime):
        return str(raw)
    if time_only:
        return raw.strftime("%H:%M:%S")
    month = _MONTHS_RU.get(raw.month, str(raw.month))
    return f"{raw.day} {month} {raw.year}, {raw.strftime('%H:%M')}"


def _format_duration(start: Any, end: Any) -> str:
    """Format duration as '3 мин 11 с', '2 ч 15 мин', etc."""
    try:
        if isinstance(start, str):
            start = datetime.fromisoformat(start.replace("Z", "+00:00"))
        if isinstance(end, str):
            end = datetime.fromisoformat(end.replace("Z", "+00:00"))
        if not isinstance(start, datetime) or not isinstance(end, datetime):
            return "—"
        delta = end - start
        total_s = int(delta.total_seconds())
    except (ValueError, TypeError):
        return "—"
    if total_s < 0:
        return "—"
    days, rem = divmod(total_s, 86400)
    hours, rem = divmod(rem, 3600)
    mins, secs = divmod(rem, 60)
    parts = []
    if days:
        parts.append(f"{days} д")
    if hours:
        parts.append(f"{hours} ч")
    if mins:
        parts.append(f"{mins} мин")
    if secs or not parts:
        parts.append(f"{secs} с")
    return " ".join(parts)


def _status_ru(raw: Any) -> str:
    value = str(raw or "").strip().upper()
    return {
        "RUNNING": "Выполняется",
        "COMPLETED": "Завершён",
        "ABORTED": "Прерван",
        "FINALIZED": "Завершён",
    }.get(value, str(raw or "—"))


def _display_value(raw: Any, *, empty: str = "не указано") -> str:
    text = str(raw or "").strip()
    return text or empty


def _add_kv_table(document: Document, rows: list[tuple[str, str]]) -> None:
    """Add a clean 2-column key-value table."""
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        cell_l = table.cell(i, 0)
        cell_l.text = label
        for run in cell_l.paragraphs[0].runs:
            run.bold = True
        table.cell(i, 1).text = value


def _existing_artifact(dataset: ReportDataset, role: str, *, category: str | None = None) -> Path | None:
    for item in dataset.artifact_index:
        if str(item.get("role", "")).strip() != role:
            continue
        if category is not None and str(item.get("category", "")).strip() != category:
            continue
        path = Path(str(item.get("path", "")).strip())
        if path.exists():
            return path
    return None


def _find_table_path(dataset: ReportDataset, table_id: str) -> Path | None:
    for item in dataset.result_tables:
        if str(item.get("table_id", "")).strip() != table_id:
            continue
        path = Path(str(item.get("path", "")).strip())
        if path.exists():
            return path
    return None


def _read_csv_preview(path: Path, *, limit: int = 6) -> tuple[list[str], list[list[str]], int]:
    """Returns (header, body_rows, total_row_count)."""
    try:
        with path.open(encoding="utf-8-sig", newline="") as handle:
            reader = csv.reader(handle)
            rows = list(reader)
    except Exception:
        with path.open(encoding="utf-8", newline="") as handle:
            reader = csv.reader(handle)
            rows = list(reader)
    if not rows:
        return [], [], 0
    header = [str(item) for item in rows[0]]
    total = len(rows) - 1
    body = [[str(item) for item in row] for row in rows[1: 1 + limit]]
    return header, body, total


def _add_table_preview(document: Document, path: Path, *, title: str, limit: int = 6) -> None:
    header, rows, total = _read_csv_preview(path, limit=limit)
    shown = len(rows)
    suffix = f", показано {shown}" if total > shown else ""
    # ГОСТ: caption above table
    num = _next_table()
    cap = document.add_paragraph()
    cap.paragraph_format.first_line_indent = Cm(0)
    run = cap.add_run(f"Таблица {num} — {title} ({total:,} строк{suffix})")
    run.font.size = Pt(12)
    run.bold = True
    if not header:
        document.add_paragraph("Нет данных.")
        return
    table = document.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for index, value in enumerate(header):
        cell = table.cell(0, index)
        cell.text = value
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(12)
    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = value
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(12)


def _add_plot(document: Document, title: str, readings: list[HistoricalReading], output_path: Path,
              *, xlabel: str = "", ylabel: str = "") -> None:
    if not readings:
        document.add_paragraph(f"{title}: данные за интервал эксперимента отсутствуют.")
        return
    output_path.parent.mkdir(parents=True, exist_ok=True)
    xs = [item.timestamp for item in readings]
    ys = [item.value for item in readings]
    plt.figure(figsize=(7, 3))
    plt.plot(xs, ys)
    plt.title(title)
    if xlabel:
        plt.xlabel(xlabel)
    if ylabel:
        plt.ylabel(ylabel)
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches="tight")
    plt.close()
    document.add_picture(str(output_path), width=Mm(160))
    # ГОСТ caption below figure, centered
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    run = cap.add_run(f"Рисунок {_next_figure()} — {title}")
    run.font.size = Pt(12)


def _add_multichannel_plot(document: Document, title: str, readings: list[HistoricalReading],
                           output_path: Path, *, xlabel: str = "", ylabel: str = "") -> None:
    """Plot multiple channels with legend (not one blob)."""
    if not readings:
        document.add_paragraph(f"{title}: данные отсутствуют.")
        return
    output_path.parent.mkdir(parents=True, exist_ok=True)
    series: dict[str, list[tuple[Any, float]]] = {}
    for r in readings:
        series.setdefault(r.channel, []).append((r.timestamp, r.value))
    plt.figure(figsize=(8, 4))
    for channel, points in sorted(series.items()):
        points.sort()
        xs = [t for t, _ in points]
        ys = [v for _, v in points]
        label = channel.split(" ")[0] if " " in channel else channel
        plt.plot(xs, ys, label=label, linewidth=0.8)
    plt.title(title)
    if xlabel:
        plt.xlabel(xlabel)
    if ylabel:
        plt.ylabel(ylabel)
    if len(series) <= 12:
        plt.legend(fontsize=7, loc="best")
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches="tight")
    plt.close()
    document.add_picture(str(output_path), width=Mm(160))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.add_run(f"Рисунок {_next_figure()} — {title}").font.size = Pt(12)


def _add_archived_or_multichannel(document: Document, dataset: ReportDataset,
                                   artifact_role: str, title: str,
                                   readings: list[HistoricalReading],
                                   fallback_path: Path, **kwargs: Any) -> None:
    """Use archived plot if available, otherwise generate multi-channel."""
    archived = _existing_artifact(dataset, artifact_role, category="plot")
    if archived:
        document.add_picture(str(archived), width=Mm(160))
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.add_run(f"Рисунок {_next_figure()} — {title}").font.size = Pt(12)
    else:
        _add_multichannel_plot(document, title, readings, fallback_path, **kwargs)


def _channel_display(raw: str) -> str:
    """Keithley_1/smua/power → SMU A: мощность."""
    if "/smua/" in raw:
        suffix = raw.split("/")[-1]
        labels = {"power": "мощность", "voltage": "напряжение", "current": "ток", "resistance": "сопротивление"}
        return f"SMU A: {labels.get(suffix, suffix)}"
    if "/smub/" in raw:
        suffix = raw.split("/")[-1]
        labels = {"power": "мощность", "voltage": "напряжение", "current": "ток", "resistance": "сопротивление"}
        return f"SMU B: {labels.get(suffix, suffix)}"
    return raw


_ROLE_RU = {
    "measured_values": "Измеренные величины (CSV)",
    "experiment_data": "Данные эксперимента (Parquet)",
    "setpoint_values": "Заданные величины (CSV)",
    "run_results": "Результаты прогонов (CSV)",
    "temperature_overview": "Температура: обзор",
    "thermal_power": "Тепловая мощность",
    "pressure": "Давление",
    "summary_metadata": "Сводка эксперимента",
    "conductivity_vs_temperature": "Теплопроводность vs температура",
}

_SOURCE_TAB_RU = {
    "Keithley 2604B": "Источник мощности",
    "keithley": "Источник мощности",
    "conductivity": "Теплопроводность",
    "autosweep": "Автоизмерение",
}


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------


def render_title_page(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    experiment = dataset.metadata["experiment"]
    template = dataset.metadata["template"]
    template_name = template.get("name", template.get("id", "Эксперимент"))
    document.add_heading(f"Отчёт: {template_name}", 0)

    title = str(experiment.get("title") or experiment.get("name") or "").strip()
    if title and title != template_name and title not in ("1", ""):
        document.add_paragraph(title).italic = True

    _add_kv_table(document, [
        ("Оператор", _display_value(experiment.get("operator"))),
        ("Образец", _display_value(experiment.get("sample"))),
        ("Криостат", _display_value(experiment.get("cryostat"))),
        ("Статус", _status_ru(experiment.get("status", ""))),
        ("Начало", _format_dt(experiment.get("start_time"))),
        ("Завершение", _format_dt(experiment.get("end_time"))),
        ("Длительность", _format_duration(experiment.get("start_time"), experiment.get("end_time"))),
        ("Идентификатор", str(experiment.get("experiment_id", ""))[:12]),
    ])

    if experiment.get("notes"):
        document.add_paragraph("")
        document.add_paragraph(f"Заметки: {experiment['notes']}")


def render_experiment_metadata_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    experiment = dataset.metadata["experiment"]
    template = dataset.metadata["template"]
    summary = dataset.summary_metadata
    document.add_heading("Метаданные эксперимента", level=1)

    # Custom fields with template labels
    custom_values = experiment.get("custom_fields", {})
    custom_defs = template.get("custom_fields", [])
    if custom_values and custom_defs:
        label_map = {str(f.get("id", "")): str(f.get("label", f.get("id", ""))) for f in custom_defs}
        rows = []
        for field_id, value in custom_values.items():
            label = label_map.get(field_id, field_id)
            rows.append((label, _display_value(value)))
        if rows:
            _add_kv_table(document, rows)
    elif custom_values:
        rows = [(k, _display_value(v)) for k, v in custom_values.items()]
        _add_kv_table(document, rows)

    # Summary stats
    if summary:
        parts = []
        if "reading_count" in summary:
            parts.append(f"Измерений: {summary['reading_count']:,}")
        if "run_count" in summary:
            parts.append(f"Прогонов: {summary['run_count']}")
        if "artifact_count" in summary:
            parts.append(f"Артефактов: {summary['artifact_count']}")
        if parts:
            document.add_paragraph(" │ ".join(parts))


def render_run_timeline_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Таймлайн прогонов", level=1)
    if not dataset.run_records:
        document.add_paragraph("Прогоны не выполнялись.")
        return

    table = document.add_table(rows=1 + len(dataset.run_records), cols=5)
    table.style = "Table Grid"
    for i, hdr in enumerate(["№", "Начало", "Конец", "Источник", "Статус"]):
        table.cell(0, i).text = hdr
    for idx, item in enumerate(dataset.run_records, 1):
        source = _SOURCE_TAB_RU.get(item.get("source_tab", ""), item.get("source_tab", "—"))
        table.cell(idx, 0).text = str(idx)
        table.cell(idx, 1).text = _format_dt(item.get("started_at"), time_only=True)
        table.cell(idx, 2).text = _format_dt(item.get("finished_at"), time_only=True)
        table.cell(idx, 3).text = source
        table.cell(idx, 4).text = _status_ru(item.get("status", ""))


def render_run_parameters_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Параметры запусков", level=1)
    if not dataset.run_records:
        document.add_paragraph("Параметры запусков отсутствуют.")
        return
    for idx, item in enumerate(dataset.run_records, 1):
        params = item.get("parameters") or {}
        if not params:
            continue
        run_type = item.get("run_type", "")
        document.add_heading(f"Прогон {idx}: {run_type}", level=2)
        rows = [(str(k), _display_value(v)) for k, v in params.items()]
        _add_kv_table(document, rows)


def render_result_tables_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Итоговые результаты и таблицы", level=1)
    table_map = {
        "measured_values": "Таблица измеренных величин",
        "setpoint_values": "Таблица заданных величин",
        "run_results": "Итоговые результаты прогонов",
        "conductivity_vs_temperature": "Теплопроводность vs температура",
    }
    rendered = False
    for table_id, title in table_map.items():
        path = _find_table_path(dataset, table_id)
        if path is None:
            continue
        _add_table_preview(document, path, title=title)
        rendered = True
    if not rendered:
        document.add_paragraph("Архивные таблицы результатов не найдены.")


def render_conductivity_section(document: Document, dataset: ReportDataset, assets_dir: Path) -> None:
    document.add_heading("Теплопроводность vs температура", level=1)
    archived_plot = _existing_artifact(dataset, "conductivity_vs_temperature", category="plot")
    if archived_plot is not None:
        document.add_picture(str(archived_plot), width=Mm(160))
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.add_run(f"Рисунок {_next_figure()} — Теплопроводность vs температура").font.size = Pt(12)
        return
    path = _find_table_path(dataset, "conductivity_vs_temperature")
    if path is None:
        document.add_paragraph("График теплопроводности для этой карточки отсутствует.")
        return
    with path.open(encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        rows = list(reader)
    if not rows:
        document.add_paragraph("График теплопроводности не удалось построить: таблица пуста.")
        return
    temps = [float(item["temperature_k"]) for item in rows]
    conds = [float(item["conductance_wk"]) for item in rows]
    plot_path = assets_dir / "conductivity_vs_temperature.png"
    plt.figure(figsize=(6, 4))
    plt.plot(temps, conds, marker="o")
    plt.title("Теплопроводность vs температура")
    plt.xlabel("Температура (К)")
    plt.ylabel("Теплопроводность (Вт/К)")
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(plot_path, dpi=150, bbox_inches="tight")
    plt.close()
    document.add_picture(str(plot_path), width=Inches(6.2))
    document.add_paragraph(
        f"Диапазон: {min(temps):.1f} — {max(temps):.1f} К, "
        f"максимум: {max(conds):.3g} Вт/К"
    )


def render_artifact_manifest_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Ключевые артефакты", level=1)
    if not dataset.artifact_index:
        document.add_paragraph("Артефакты в архивной карточке отсутствуют.")
        return

    by_category: dict[str, list[dict]] = defaultdict(list)
    for item in dataset.artifact_index:
        cat = str(item.get("category", "other")).strip() or "other"
        by_category[cat].append(item)

    cat_names = {"table": "Таблицы", "plot": "Графики", "summary": "Сводки", "other": "Прочее"}
    for cat, items in by_category.items():
        document.add_paragraph(cat_names.get(cat, cat), style="List Bullet")
        for item in items:
            role = str(item.get("role", "")).strip()
            name_ru = _ROLE_RU.get(role, role)
            path = Path(str(item.get("path", "")))
            document.add_paragraph(f"    {name_ru} — {path.name}")


def render_cooldown_section(document: Document, dataset: ReportDataset, assets_dir: Path) -> None:
    document.add_heading("Охлаждение", level=1)
    temp_readings = [item for item in dataset.readings if item.unit == "K"]
    _add_archived_or_multichannel(
        document, dataset, "temperature_overview", "Температура каналов",
        temp_readings, assets_dir / "cooldown_temperature.png",
        xlabel="Время", ylabel="Температура (К)",
    )
    if temp_readings:
        t_init = temp_readings[0].value
        t_final = temp_readings[-1].value
        duration = _format_duration(temp_readings[0].timestamp, temp_readings[-1].timestamp)
        _add_kv_table(document, [
            ("Начальная температура", f"{t_init:.2f} К"),
            ("Конечная температура", f"{t_final:.2f} К"),
            ("Время охлаждения", duration),
            ("Число точек", f"{len(temp_readings):,}"),
        ])

        # Check target temperature from custom fields
        custom = dataset.metadata.get("experiment", {}).get("custom_fields", {})
        target = custom.get("target_temperature")
        if target:
            try:
                target_k = float(target)
                reached = t_final <= target_k * 1.05  # 5% tolerance
                document.add_paragraph(
                    f"Целевая: {target_k:.1f} К — {'достигнута ✓' if reached else 'не достигнута'}"
                )
            except (ValueError, TypeError):
                pass


def render_thermal_section(document: Document, dataset: ReportDataset, assets_dir: Path) -> None:
    document.add_heading("Тепловая нагрузка", level=1)
    power_readings = [item for item in dataset.readings if item.channel.endswith("/power")]
    _add_archived_or_multichannel(
        document, dataset, "thermal_power", "Мощность Keithley",
        power_readings, assets_dir / "thermal_power.png",
        xlabel="Время", ylabel="Мощность (Вт)",
    )
    if power_readings:
        by_channel: dict[str, list[float]] = defaultdict(list)
        for item in power_readings:
            by_channel[item.channel].append(item.value)

        table = document.add_table(rows=1 + len(by_channel), cols=3)
        table.style = "Table Grid"
        for i, hdr in enumerate(["Канал", "Средняя мощность", "Макс. мощность"]):
            table.cell(0, i).text = hdr
        for idx, (channel, values) in enumerate(sorted(by_channel.items()), 1):
            table.cell(idx, 0).text = _channel_display(channel)
            table.cell(idx, 1).text = f"{mean(values):.4g} Вт"
            table.cell(idx, 2).text = f"{max(values):.4g} Вт"


def render_pressure_section(document: Document, dataset: ReportDataset, assets_dir: Path) -> None:
    document.add_heading("Давление", level=1)
    pressure = [
        item for item in dataset.readings if "pressure" in item.channel.lower() or item.unit.lower() in {"mbar", "pa"}
    ]
    _add_archived_or_multichannel(
        document, dataset, "pressure", "Давление",
        pressure, assets_dir / "pressure.png",
        xlabel="Время", ylabel="Давление (мбар)",
    )
    if pressure:
        vals = [p.value for p in pressure if math.isfinite(p.value) and p.value > 0]
        if vals:
            _add_kv_table(document, [
                ("Число точек", f"{len(pressure):,}"),
                ("Последнее значение", f"{pressure[-1].value:.3e} {pressure[-1].unit}"),
                ("Минимум", f"{min(vals):.3e} мбар"),
                ("Максимум", f"{max(vals):.3e} мбар"),
            ])


def render_operator_log_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Служебный лог", level=1)
    if not dataset.operator_log:
        document.add_paragraph("Записи служебного лога за интервал эксперимента отсутствуют.")
        return
    entries = dataset.operator_log
    total = len(entries)
    if total > 30:
        show = list(entries[:15]) + list(entries[-5:])
        skipped = total - 20
    else:
        show = entries
        skipped = 0

    for item in show[:15]:
        who = item.author or item.source or "система"
        tag_suffix = f" [{', '.join(item.tags)}]" if item.tags else ""
        document.add_paragraph(
            f"{_format_dt(item.timestamp, time_only=True)} │ {who}: {item.message}{tag_suffix}",
            style="List Bullet",
        )
    if skipped > 0:
        document.add_paragraph(f"... ещё {skipped} записей ...")
        for item in show[15:]:
            who = item.author or item.source or "система"
            tag_suffix = f" [{', '.join(item.tags)}]" if item.tags else ""
            document.add_paragraph(
                f"{_format_dt(item.timestamp, time_only=True)} │ {who}: {item.message}{tag_suffix}",
                style="List Bullet",
            )


def render_alarms_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Алармы", level=1)
    if not dataset.alarm_readings:
        document.add_paragraph("Алармов не зафиксировано ✓")
        return
    for item in dataset.alarm_readings[-20:]:
        document.add_paragraph(
            f"{_format_dt(item.timestamp, time_only=True)} │ "
            f"{_channel_display(item.channel)} = {item.value:g} {item.unit} [{item.status}]",
            style="List Bullet",
        )


def render_config_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Снимок конфигурации", level=1)
    config_snapshot = dataset.metadata["experiment"].get("config_snapshot", {})
    if not config_snapshot:
        document.add_paragraph("Снимок конфигурации для этого эксперимента не сохранён.")
        return

    # Instruments table
    instruments = config_snapshot.get("instruments", [])
    if instruments:
        document.add_heading("Приборы", level=2)
        table = document.add_table(rows=1 + len(instruments), cols=3)
        table.style = "Table Grid"
        for i, hdr in enumerate(["Прибор", "Интерфейс", "Адрес"]):
            table.cell(0, i).text = hdr
        for idx, instr in enumerate(instruments, 1):
            name = str(instr.get("name", instr.get("type", "—")))
            itype = str(instr.get("type", "—"))
            resource = str(instr.get("resource", instr.get("resource_str", "—")))
            iface = "GPIB" if "GPIB" in resource.upper() else "USB-TMC" if "USB" in resource.upper() else "RS-232" if "COM" in resource.upper() else "—"
            table.cell(idx, 0).text = name
            table.cell(idx, 1).text = iface
            table.cell(idx, 2).text = resource

    # Safety limits
    safety = config_snapshot.get("safety", {})
    if safety:
        document.add_heading("Лимиты безопасности", level=2)
        rows = []
        if "max_power_w" in safety:
            rows.append(("Макс. мощность", f"{safety['max_power_w']} Вт"))
        if "max_voltage_v" in safety:
            rows.append(("Макс. напряжение", f"{safety['max_voltage_v']} В"))
        if "max_current_a" in safety:
            rows.append(("Макс. ток", f"{safety['max_current_a']} А"))
        if "max_dT_dt_K_per_min" in safety:
            rows.append(("Макс. dT/dt", f"{safety['max_dT_dt_K_per_min']} К/мин"))
        if rows:
            _add_kv_table(document, rows)


def render_operator_comments_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Комментарии оператора", level=1)
    document.add_paragraph("Заполнить после автоматической генерации отчёта.")
    document.add_paragraph("")
    document.add_paragraph("")


def render_operator_interpretation_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Интерпретация результатов", level=1)
    document.add_paragraph("Сюда оператор добавляет интерпретацию, выводы и замечания по качеству данных.")
    document.add_paragraph("")
    document.add_paragraph("")


def render_operator_photos_section(document: Document, dataset: ReportDataset, _assets_dir: Path) -> None:
    document.add_heading("Фотографии и внешние изображения", level=1)
    document.add_paragraph("В этот раздел можно вставить фотографии образца, оснастки и внешние иллюстрации.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Изображение 1"
    table.cell(0, 1).text = "Описание"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

SECTION_REGISTRY: dict[str, SectionRenderer] = {
    "title_page": render_title_page,
    "experiment_metadata_section": render_experiment_metadata_section,
    "run_timeline_section": render_run_timeline_section,
    "run_parameters_section": render_run_parameters_section,
    "result_tables_section": render_result_tables_section,
    "conductivity_section": render_conductivity_section,
    "artifact_manifest_section": render_artifact_manifest_section,
    "cooldown_section": render_cooldown_section,
    "thermal_section": render_thermal_section,
    "pressure_section": render_pressure_section,
    "operator_log_section": render_operator_log_section,
    "alarms_section": render_alarms_section,
    "config_section": render_config_section,
    "operator_comments_section": render_operator_comments_section,
    "operator_interpretation_section": render_operator_interpretation_section,
    "operator_photos_section": render_operator_photos_section,
}
