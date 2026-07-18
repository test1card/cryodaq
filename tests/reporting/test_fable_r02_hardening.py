from __future__ import annotations

import json
from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from cryodaq.reporting.data import HistoricalReading, ReportDataExtractor, ReportDataset
from cryodaq.reporting.sections import render_cooldown_section, render_thermal_section


def _document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_dataset_with_missing_end_time_uses_a_bounded_current_fallback(
    tmp_path: Path,
    monkeypatch,
) -> None:
    metadata_path = tmp_path / "metadata.json"
    metadata_path.write_text(
        json.dumps(
            {
                "experiment": {
                    "experiment_id": "exp-open",
                    "start_time": "2026-07-17T12:00:00+00:00",
                },
                "artifact_index": [],
                "result_tables": [],
            }
        ),
        encoding="utf-8",
    )
    extractor = ReportDataExtractor(tmp_path)
    observed: dict[str, datetime] = {}

    def load_readings(start: datetime, end: datetime) -> list:
        observed["start"] = start
        observed["end"] = end
        return []

    monkeypatch.setattr(extractor, "_load_readings", load_readings)
    monkeypatch.setattr(extractor, "_load_operator_log", lambda *_args: [])

    before = datetime.now(UTC)
    dataset = extractor.load_dataset(metadata_path)
    after = datetime.now(UTC)

    assert dataset.readings == []
    assert observed["start"] == datetime(2026, 7, 17, 12, tzinfo=UTC)
    assert before <= observed["end"] <= after


def test_nan_temperature_and_power_render_explicit_no_data(tmp_path: Path) -> None:
    timestamp = datetime(2026, 7, 17, 12, tzinfo=UTC)
    dataset = ReportDataset(
        metadata={"experiment": {}, "template": {}},
        readings=[
            HistoricalReading(timestamp, "probe", "T1", float("nan"), "K", "sensor_error"),
            HistoricalReading(timestamp, "source", "smua/power", float("nan"), "W", "sensor_error"),
        ],
    )
    document = Document()

    render_cooldown_section(document, dataset, tmp_path)
    render_thermal_section(document, dataset, tmp_path)

    text = _document_text(document).lower()
    assert "nan" not in text
    assert text.count("данные отсутствуют") >= 2
