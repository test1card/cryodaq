from __future__ import annotations

import json
import subprocess
from datetime import UTC, datetime
from pathlib import Path

import pytest
import yaml
from docx import Document

from cryodaq.core.experiment import ExperimentManager
from cryodaq.drivers.base import ChannelStatus, Reading
from cryodaq.report_process import ReportProcessRunner
from cryodaq.report_state import ReportContractError, load_current_manifest
from cryodaq.reporting.generator import ReportGenerator
from cryodaq.storage.sqlite_writer import SQLiteWriter


@pytest.fixture()
def instruments_yaml(tmp_path: Path) -> Path:
    path = tmp_path / "instruments.yaml"
    path.write_text(
        yaml.dump({"instruments": [{"name": "k1", "type": "keithley_2604b", "resource": "MOCK"}]}),
        encoding="utf-8",
    )
    return path


@pytest.fixture()
def templates_dir(tmp_path: Path) -> Path:
    root = tmp_path / "experiment_templates"
    root.mkdir()
    (root / "thermal_conductivity.yaml").write_text(
        yaml.dump(
            {
                "id": "thermal_conductivity",
                "name": "Thermal Conductivity",
                "sections": ["setup", "sample"],
                "report_enabled": True,
                "report_sections": [
                    "thermal_section",
                    "pressure_section",
                    "config_section",
                ],
            }
        ),
        encoding="utf-8",
    )
    (root / "cooldown_test.yaml").write_text(
        yaml.dump(
            {
                "id": "cooldown_test",
                "name": "Cooldown Test",
                "sections": ["setup", "cooldown_path"],
                "report_enabled": True,
                "report_sections": [
                    "cooldown_section",
                    "pressure_section",
                    "alarms_section",
                ],
            }
        ),
        encoding="utf-8",
    )
    (root / "debug_checkout.yaml").write_text(
        yaml.dump(
            {
                "id": "debug_checkout",
                "name": "Debug Checkout",
                "sections": ["setup", "checks"],
                "report_enabled": False,
                "report_sections": ["config_section"],
            }
        ),
        encoding="utf-8",
    )
    return root


@pytest.fixture()
def manager(tmp_path: Path, instruments_yaml: Path, templates_dir: Path) -> ExperimentManager:
    return ExperimentManager(tmp_path, instruments_yaml, templates_dir=templates_dir)


def _reading(channel: str, value: float, unit: str, ts: datetime) -> Reading:
    return Reading(
        timestamp=ts,
        instrument_id="k1",
        channel=channel,
        value=value,
        unit=unit,
        status=ChannelStatus.OK,
    )


async def _seed_experiment_data(tmp_path: Path, experiment_id: str) -> None:
    writer = SQLiteWriter(tmp_path)
    ts = datetime(2026, 3, 16, 12, 0, tzinfo=UTC)
    writer._write_batch(
        [
            _reading("K1/smua/power", 1.2, "W", ts),
            _reading("K1/smub/power", 0.8, "W", ts),
            _reading("P_MAIN/pressure", 2.1e-4, "mbar", ts),
            _reading("T_STAGE", 4.3, "K", ts),
            _reading("alarm/high_pressure", 1.0, "", ts),
        ]
    )
    await writer.append_operator_log(
        message="Report marker",
        author="ivanov",
        source="gui",
        experiment_id=experiment_id,
        timestamp=ts,
    )
    await writer.stop()


def _doc_text(path: Path) -> str:
    document = Document(path)
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def test_direct_generator_rejects_experiment_traversal(tmp_path: Path) -> None:
    with pytest.raises(ReportContractError):
        ReportGenerator(tmp_path).generate("../outside")


async def test_direct_generator_rejects_missing_metadata_artifact(
    manager: ExperimentManager,
    tmp_path: Path,
) -> None:
    exp_id = manager.start_experiment(
        name="Missing artifact",
        operator="Operator",
        template_id="thermal_conductivity",
        start_time="2026-03-16T12:00:00+00:00",
    )
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")
    metadata_path = tmp_path / "experiments" / exp_id / "metadata.json"
    payload = json.loads(metadata_path.read_text(encoding="utf-8"))
    payload.setdefault("artifact_index", []).append(
        {"role": "missing", "path": str(metadata_path.parent / "missing.csv")}
    )
    metadata_path.write_text(json.dumps(payload), encoding="utf-8")

    with pytest.raises(ReportContractError, match="must exist"):
        ReportGenerator(tmp_path).generate(exp_id)


async def test_direct_generator_rejects_symlinked_archive_input(
    manager: ExperimentManager,
    tmp_path: Path,
) -> None:
    exp_id = manager.start_experiment(
        name="Symlink artifact",
        operator="Operator",
        template_id="thermal_conductivity",
        start_time="2026-03-16T12:00:00+00:00",
    )
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")
    experiment_root = tmp_path / "experiments" / exp_id
    outside = tmp_path / "outside.csv"
    outside.write_text("a,b\n1,2\n", encoding="utf-8")
    link = experiment_root / "archive" / "tables" / "linked.csv"
    link.symlink_to(outside)
    metadata_path = experiment_root / "metadata.json"
    payload = json.loads(metadata_path.read_text(encoding="utf-8"))
    payload.setdefault("result_tables", []).append({"table_id": "linked", "path": str(link)})
    metadata_path.write_text(json.dumps(payload), encoding="utf-8")

    with pytest.raises(ReportContractError, match="symlink"):
        ReportGenerator(tmp_path).generate(exp_id)


async def test_ephemeral_child_preserves_manual_schema_and_selects_generation(
    manager: ExperimentManager,
    tmp_path: Path,
) -> None:
    exp_id = manager.start_experiment(
        name="Child report",
        title="Child report",
        operator="Operator",
        template_id="thermal_conductivity",
        start_time="2026-03-16T12:00:00+00:00",
    )
    await _seed_experiment_data(tmp_path, exp_id)
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")

    report = ReportProcessRunner(tmp_path, timeout_s=20).generate_experiment(exp_id)

    assert set(report) == {
        "docx_path",
        "pdf_path",
        "assets_dir",
        "sections",
        "skipped",
        "reason",
    }
    assert isinstance(report["docx_path"], str)
    assert report["pdf_path"] is None or isinstance(report["pdf_path"], str)
    assert isinstance(report["assets_dir"], str)
    assert isinstance(report["sections"], list)
    assert type(report["skipped"]) is bool
    assert isinstance(report["reason"], str)
    assert Path(report["docx_path"]).is_file()  # noqa: ASYNC240 - child completed
    manifest = load_current_manifest(tmp_path / "experiments" / exp_id)
    assert manifest is not None
    report_path = Path(report["docx_path"])
    assert report_path.relative_to(tmp_path / "experiments" / exp_id) == (
        Path("reports") / "generations" / manifest["generation_id"] / "report_editable.docx"
    )
    archive = manager.get_archive_item(exp_id)
    assert archive is not None
    assert archive.docx_path == Path(report["docx_path"])


async def test_ephemeral_child_preserves_disabled_report_types(
    manager: ExperimentManager,
    tmp_path: Path,
) -> None:
    exp_id = manager.start_experiment(
        name="Disabled child report",
        title="Disabled child report",
        operator="Operator",
        template_id="debug_checkout",
        start_time="2026-03-16T12:00:00+00:00",
    )
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")

    report = ReportProcessRunner(tmp_path, timeout_s=20).generate_experiment(exp_id)

    assert report["skipped"] is True
    assert report["pdf_path"] is None
    assert report["sections"] == []
    assert isinstance(report["docx_path"], str)
    assert isinstance(report["assets_dir"], str)


async def test_report_generation_uses_new_output_names_and_sections(manager: ExperimentManager, tmp_path: Path) -> None:
    exp_id = manager.start_experiment(
        name="Lambda",
        title="Lambda",
        operator="Ivanov",
        template_id="thermal_conductivity",
        sample="Cu sample",
        notes="Card note",
        start_time="2026-03-16T12:00:00+00:00",
    )
    await _seed_experiment_data(tmp_path, exp_id)
    manager.attach_run_record(
        experiment_id=exp_id,
        source_tab="autosweep",
        source_module="autosweep_panel",
        run_type="autosweep",
        status="COMPLETED",
        started_at="2026-03-16T12:01:00+00:00",
        finished_at="2026-03-16T12:02:00+00:00",
        source_run_id="autosweep-1",
        parameters={"power_start_w": 0.1},
        result_summary={"point_count": 3},
        artifact_paths=[],
    )
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")

    result = ReportGenerator(tmp_path).generate(exp_id)

    assert result.skipped is False
    assert result.docx_path.name == "report_editable.docx"
    assert (tmp_path / "experiments" / exp_id / "reports" / "report_raw.docx").exists()
    assert "run_timeline_section" in result.sections
    assert "run_parameters_section" in result.sections
    assert "result_tables_section" in result.sections
    assert "conductivity_section" in result.sections
    assert "artifact_manifest_section" in result.sections
    assert "operator_comments_section" in result.sections
    assert "operator_interpretation_section" in result.sections
    assert "operator_photos_section" in result.sections

    text = _doc_text(result.docx_path)
    assert "Ivanov" in text
    assert "Cu sample" in text
    assert "Таймлайн прогонов" in text
    assert "Параметры запусков" in text
    assert "Итоговые результаты и таблицы" in text
    assert "Ключевые артефакты" in text
    assert "Комментарии оператора" in text
    assert "Интерпретация результатов" in text
    assert "Фотографии и внешние изображения" in text


async def test_report_generation_for_cooldown_template_uses_archive_tables(
    manager: ExperimentManager, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("CRYODAQ_ALLOW_BROKEN_SQLITE", "1")
    exp_id = manager.start_experiment(
        name="Cooldown",
        title="Cooldown",
        operator="Petrov",
        template_id="cooldown_test",
        start_time="2026-03-16T12:00:00+00:00",
    )
    await _seed_experiment_data(tmp_path, exp_id)
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")

    # Remove the live DB so that generate() MUST fall back to the archive CSV.
    # If the generator silently skips the archive and reads only from the live
    # DB, the temperature/pressure data would be absent from the report.
    live_db = tmp_path / "data_2026-03-16.db"
    if live_db.exists():
        live_db.unlink()

    result = ReportGenerator(tmp_path).generate(exp_id)

    archive_csv = tmp_path / "experiments" / exp_id / "archive" / "tables" / "measured_values.csv"
    assert archive_csv.exists(), "finalize must produce measured_values.csv"

    text = _doc_text(result.docx_path)
    assert "Охлаждение" in text
    assert "Тревоги" in text
    # "Таблица измеренных величин" only appears if cooldown_test template
    # is extended to include result_tables_section; assert the seeded
    # temperature data (T_STAGE = 4.3 K) appears via the cooldown kv_table
    # — proving readings were loaded from the archive CSV, not the deleted DB.
    assert "4.30 К" in text, "Seeded T_STAGE=4.3 K must appear in report — archive CSV not read"


async def test_report_disabled_template_is_respected(manager: ExperimentManager, tmp_path: Path) -> None:
    exp_id = manager.start_experiment(
        name="Checkout",
        title="Checkout",
        operator="Sidorov",
        template_id="debug_checkout",
        start_time="2026-03-16T12:00:00+00:00",
    )
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")

    result = ReportGenerator(tmp_path).generate(exp_id)

    assert result.skipped is True
    assert result.reason == "Формирование отчёта отключено шаблоном."
    assert result.docx_path.exists() is False


async def test_report_artifact_folder_contains_named_outputs(manager: ExperimentManager, tmp_path: Path) -> None:
    exp_id = manager.start_experiment(
        name="Artifact Check",
        title="Artifact Check",
        operator="Operator",
        template_id="cooldown_test",
        start_time="2026-03-16T12:00:00+00:00",
    )
    await _seed_experiment_data(tmp_path, exp_id)
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")

    result = ReportGenerator(tmp_path).generate(exp_id)
    report_dir = tmp_path / "experiments" / exp_id / "reports"

    assert report_dir.exists()
    assert result.docx_path == report_dir / "report_editable.docx"
    assert result.assets_dir.exists()
    assert (report_dir / "report_raw.docx").exists()


async def test_report_generation_graceful_without_pdf_tooling(
    manager: ExperimentManager,
    tmp_path: Path,
    monkeypatch,
    caplog: pytest.LogCaptureFixture,
) -> None:
    exp_id = manager.start_experiment(
        name="No PDF",
        title="No PDF",
        operator="Operator",
        template_id="cooldown_test",
        start_time="2026-03-16T12:00:00+00:00",
    )
    await _seed_experiment_data(tmp_path, exp_id)
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")

    monkeypatch.setattr("cryodaq.reporting.generator.shutil.which", lambda _name: None)
    with caplog.at_level("WARNING", logger="cryodaq.reporting.generator"):
        result = ReportGenerator(tmp_path).generate(exp_id)

    assert result.docx_path.exists()
    assert result.docx_path.name == "report_editable.docx"
    assert result.pdf_path is None
    # Degradation must be LOUD: a WARNING naming the consequence + the remedy.
    assert any(
        rec.levelname == "WARNING" and "PDF не создан" in rec.message and "LibreOffice" in rec.message
        for rec in caplog.records
    ), "missing-soffice degradation must log a WARNING naming PDF loss + LibreOffice remedy"


async def test_report_generation_can_use_archived_measured_values_without_live_db(
    manager: ExperimentManager,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("CRYODAQ_ALLOW_BROKEN_SQLITE", "1")
    exp_id = manager.start_experiment(
        name="Archive source",
        title="Archive source",
        operator="Operator",
        template_id="cooldown_test",
        start_time="2026-03-16T12:00:00+00:00",
    )
    writer = SQLiteWriter(tmp_path)
    ts = datetime(2026, 3, 16, 12, 1, tzinfo=UTC)
    writer._write_batch(
        [
            _reading("K1/smua/power", 1.2, "W", ts),
            _reading("P_MAIN/pressure", 2.1e-4, "mbar", ts),
            _reading("T_STAGE", 4.3, "K", ts),
        ]
    )
    if writer._conn is not None:
        writer._conn.close()
        writer._conn = None
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")

    # Delete live DB — generator must fall back to the archive CSV exclusively.
    db_path = tmp_path / "data_2026-03-16.db"
    if db_path.exists():
        db_path.unlink()

    result = ReportGenerator(tmp_path).generate(exp_id)

    assert result.docx_path.exists()
    archive_csv = tmp_path / "experiments" / exp_id / "archive" / "tables" / "measured_values.csv"
    assert archive_csv.exists(), "finalize must produce measured_values.csv"

    text = _doc_text(result.docx_path)
    # Seeded T_STAGE=4.3 K must surface in the cooldown kv_table ("4.30 К"),
    # proving the extractor loaded readings from the archive CSV not the live DB.
    assert "4.30 К" in text, "Seeded T_STAGE=4.3 K not found in report — archive CSV path broken"
    # Seeded pressure reading channel name must appear in the archive CSV.
    csv_text = archive_csv.read_text(encoding="utf-8")
    assert "K1/smua/power" in csv_text, "K1/smua/power not found in archive CSV"
    assert "P_MAIN/pressure" in csv_text, "P_MAIN/pressure not found in archive CSV"
    assert "T_STAGE" in csv_text, "T_STAGE not found in archive CSV"


async def test_report_generation_graceful_on_soffice_timeout(
    manager: ExperimentManager,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A hung soffice (TimeoutExpired) must degrade to docx-only, not raise.

    Falls back exactly like the missing-soffice path: pdf_path is None and
    the editable docx is still produced. No exception bubbles to the caller.
    """
    monkeypatch.setenv("CRYODAQ_ALLOW_BROKEN_SQLITE", "1")
    exp_id = manager.start_experiment(
        name="Timeout PDF",
        title="Timeout PDF",
        operator="Operator",
        template_id="cooldown_test",
        start_time="2026-03-16T12:00:00+00:00",
    )
    await _seed_experiment_data(tmp_path, exp_id)
    manager.finalize_experiment(exp_id, end_time="2026-03-16T12:05:00+00:00")

    # Pretend soffice exists, but make the subprocess hang past its timeout.
    monkeypatch.setattr("cryodaq.reporting.generator.shutil.which", lambda _name: "/usr/bin/soffice")

    def _raise_timeout(*_args, **_kwargs):
        raise subprocess.TimeoutExpired(cmd="soffice", timeout=120)

    monkeypatch.setattr("cryodaq.reporting.generator.subprocess.run", _raise_timeout)

    result = ReportGenerator(tmp_path).generate(exp_id)

    assert result.docx_path.exists()
    assert result.docx_path.name == "report_editable.docx"
    assert result.pdf_path is None


def test_deadline_conversion_reserves_frozen_commit_tail(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import cryodaq.reporting.generator as generator_module

    waits: list[float] = []
    terminated: list[int] = []

    class HungSoffice:
        pid = 42

        def wait(self, timeout: float) -> int:
            waits.append(timeout)
            if len(waits) == 1:
                raise subprocess.TimeoutExpired(cmd="soffice", timeout=timeout)
            return 0

    monkeypatch.setattr(generator_module.shutil, "which", lambda _name: "soffice")
    monkeypatch.setattr(generator_module.subprocess, "Popen", lambda *_args, **_kwargs: HungSoffice())
    monkeypatch.setattr(generator_module.time, "time", lambda: 100.0)
    monkeypatch.setattr(generator_module, "terminate_descendant_tree", terminated.append)

    result = ReportGenerator(tmp_path)._try_convert_pdf(
        tmp_path / "report.docx",
        tmp_path / "report.pdf",
        deadline_epoch=115.0,
    )

    assert result is None
    assert waits == [2.0, 2.0]
    assert terminated == [42]


@pytest.mark.parametrize(
    "tree_error",
    [
        subprocess.TimeoutExpired(cmd="taskkill", timeout=2.0),
        OSError("taskkill failed"),
    ],
)
def test_deadline_cleanup_falls_back_to_direct_leader_kill(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    tree_error: BaseException,
) -> None:
    import cryodaq.reporting.generator as generator_module

    waits: list[float] = []
    kills: list[int] = []

    class HungSoffice:
        pid = 45

        def wait(self, timeout: float) -> int:
            waits.append(timeout)
            if len(waits) == 1:
                raise subprocess.TimeoutExpired(cmd="soffice", timeout=timeout)
            return 0

        def kill(self) -> None:
            kills.append(self.pid)

    def fail_tree_cleanup(_pid: int) -> None:
        raise tree_error

    monkeypatch.setattr(generator_module.shutil, "which", lambda _name: "soffice")
    monkeypatch.setattr(generator_module.subprocess, "Popen", lambda *_args, **_kwargs: HungSoffice())
    monkeypatch.setattr(generator_module.time, "time", lambda: 100.0)
    monkeypatch.setattr(generator_module, "terminate_descendant_tree", fail_tree_cleanup)

    result = ReportGenerator(tmp_path)._try_convert_pdf(
        tmp_path / "report.docx",
        tmp_path / "report.pdf",
        deadline_epoch=115.0,
    )

    assert result is None
    assert waits == [2.0, 2.0]
    assert kills == [45]


def test_deadline_conversion_does_not_enter_commit_reserve(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import cryodaq.reporting.generator as generator_module

    monkeypatch.setattr(generator_module.shutil, "which", lambda _name: "soffice")
    monkeypatch.setattr(generator_module.time, "time", lambda: 100.0)
    monkeypatch.setattr(
        generator_module.subprocess,
        "Popen",
        lambda *_args, **_kwargs: pytest.fail("soffice must not start inside commit reserve"),
    )

    result = ReportGenerator(tmp_path)._try_convert_pdf(
        tmp_path / "report.docx",
        tmp_path / "report.pdf",
        deadline_epoch=113.0,
    )

    assert result is None


def test_deadline_conversion_recomputes_budget_after_process_start(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import cryodaq.reporting.generator as generator_module

    waits: list[float] = []
    clock = iter((100.0, 102.0))

    class Soffice:
        pid = 43

        def wait(self, timeout: float) -> int:
            waits.append(timeout)
            return 0

    monkeypatch.setattr(generator_module.shutil, "which", lambda _name: "soffice")
    monkeypatch.setattr(generator_module.subprocess, "Popen", lambda *_args, **_kwargs: Soffice())
    monkeypatch.setattr(generator_module.time, "time", lambda: next(clock, 102.0))

    result = ReportGenerator(tmp_path)._try_convert_pdf(
        tmp_path / "report.docx",
        tmp_path / "report.pdf",
        deadline_epoch=120.0,
    )

    assert result is None
    assert waits == [5.0]


def test_deadline_conversion_settles_process_when_start_consumes_budget(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import cryodaq.reporting.generator as generator_module

    waits: list[float] = []
    terminated: list[int] = []
    clock = iter((100.0, 103.0))

    class Soffice:
        pid = 44

        def wait(self, timeout: float) -> int:
            waits.append(timeout)
            return 0

    monkeypatch.setattr(generator_module.shutil, "which", lambda _name: "soffice")
    monkeypatch.setattr(generator_module.subprocess, "Popen", lambda *_args, **_kwargs: Soffice())
    monkeypatch.setattr(generator_module.time, "time", lambda: next(clock, 103.0))
    monkeypatch.setattr(generator_module, "terminate_descendant_tree", terminated.append)

    result = ReportGenerator(tmp_path)._try_convert_pdf(
        tmp_path / "report.docx",
        tmp_path / "report.pdf",
        deadline_epoch=115.0,
    )

    assert result is None
    assert terminated == [44]
    assert waits == [2.0]


def test_gemma_intro_with_control_char_renders_without_raising(tmp_path: Path) -> None:
    """LLM intro containing a C0 control char must be xml_safe-wrapped (no crash)."""
    from docx import Document

    from cryodaq.reporting.generator import ReportGenerator

    document = Document()
    # A bell char (\x07) is illegal in XML 1.0 and would raise inside python-docx
    # if passed straight through to add_paragraph.
    ReportGenerator._render_gemma_annotation(document, "Введение.\x07\nВторой абзац.")

    text = "\n".join(p.text for p in document.paragraphs if p.text)
    assert "Аннотация" in text
    assert "Введение." in text
    assert "Второй абзац." in text
    assert "\x07" not in text


def test_service_log_empty_state_is_russian(tmp_path: Path) -> None:
    from cryodaq.reporting.data import ReportDataset
    from cryodaq.reporting.sections import render_operator_log_section

    document = Document()
    dataset = ReportDataset(metadata={"experiment": {}, "template": {}}, operator_log=[])

    render_operator_log_section(document, dataset, tmp_path)

    text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    assert "Служебный лог" in text
    assert "Записи служебного лога за интервал эксперимента отсутствуют." in text
