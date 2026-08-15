"""Unavailable readings must render as "—", never as a confident "nan".

NaN-доктрина (reporting/data.py::_parse_archived_value): a non-finite value
means "no reading". Two DOCX surfaces printed it verbatim — the pressure
summary's "Последнее значение" and the alarm list's value column — so an
operator read `nan mbar` where the sensor had simply dropped out.
"""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

import docx
import pytest

from cryodaq.reporting.data import HistoricalReading, ReportDataset
from cryodaq.reporting.sections import (
    render_alarms_section,
    render_experiment_metadata_section,
    render_pressure_section,
)

NAN = float("nan")
_TS = datetime(2026, 3, 14, 12, 0, 0, tzinfo=UTC)


def _document_text(document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


@pytest.fixture
def document():
    return docx.Document()


def test_pressure_last_value_unavailable_renders_dash(document, tmp_path: Path) -> None:
    """Last pressure sample is NaN — the summary must not print "nan"."""
    dataset = ReportDataset(
        metadata={"experiment": {}},
        readings=[
            HistoricalReading(_TS, "mks", "vacuum/pressure", 1.2e-5, "mbar", "ok"),
            HistoricalReading(_TS, "mks", "vacuum/pressure", NAN, "mbar", "stale"),
        ],
    )

    render_pressure_section(document, dataset, tmp_path)
    text = _document_text(document)

    assert "nan" not in text.lower(), f"DOCX rendered a non-finite reading verbatim:\n{text}"
    assert "данные отсутствуют" in text
    # Legacy readings have no descriptor authority and are excluded.
    assert "1.200e-05" not in text


def test_pressure_last_value_finite_still_rendered(document, tmp_path: Path) -> None:
    """Guard: the dash must not swallow a real reading."""
    dataset = ReportDataset(
        metadata={"experiment": {}},
        readings=[HistoricalReading(_TS, "mks", "vacuum/pressure", 3.4e-6, "mbar", "ok")],
    )

    render_pressure_section(document, dataset, tmp_path)
    text = _document_text(document)

    assert "3.400e-06" not in text


def test_alarm_reading_unavailable_renders_dash(document, tmp_path: Path) -> None:
    """An alarm whose triggering value is unavailable must not print "nan"."""
    dataset = ReportDataset(
        metadata={"experiment": {}},
        alarm_readings=[HistoricalReading(_TS, "eng", "alarm/vac_high", NAN, "mbar", "TRIGGERED")],
    )

    render_alarms_section(document, dataset, tmp_path)
    text = _document_text(document)

    assert "nan" not in text.lower(), f"DOCX rendered a non-finite alarm value:\n{text}"
    assert "—" in text
    # The alarm itself is still reported — only the number is withheld.
    assert "TRIGGERED" in text


def test_alarm_reading_finite_still_rendered(document, tmp_path: Path) -> None:
    dataset = ReportDataset(
        metadata={"experiment": {}},
        alarm_readings=[HistoricalReading(_TS, "eng", "alarm/vac_high", 12.5, "mbar", "TRIGGERED")],
    )

    render_alarms_section(document, dataset, tmp_path)
    text = _document_text(document)

    assert "12.5 mbar" in text


def test_summary_counts_are_derived_from_loaded_report_collections(document, tmp_path: Path) -> None:
    """Keys are the producer's: `_build_archive_snapshot` writes these three and no others."""
    dataset = ReportDataset(
        metadata={"experiment": {}, "template": {}},
        readings=[HistoricalReading(_TS, "meter", "input", 1.0, "V", "ok")],
        run_records=[{"status": "COMPLETED"}],
        artifact_index=[{"role": "measured_values"}, {"role": "summary_metadata"}],
        summary_metadata={"measured_value_rows": 0, "run_record_count": 0, "artifact_count": 0},
    )

    render_experiment_metadata_section(document, dataset, tmp_path)
    text = _document_text(document)

    assert "Измерений: 1" in text, "report rendered an unchecked measured_value_rows claim"
    assert "Прогонов: 1" in text, "report rendered an unchecked run_record_count claim"
    # The whole loaded index is shown; the stale record is shown beside it, not instead of it.
    assert "Артефактов: 2 (в записи: 0)" in text, "report rendered an unchecked artifact_count claim"


def test_summary_artifact_count_compares_the_membership_the_archive_counted(document, tmp_path: Path) -> None:
    """A consistent archive must not be reported as stale.

    `_build_archive_snapshot` records `artifact_count` before appending its own
    summary_metadata artifact, so the loaded index of a fully consistent archive is exactly
    one longer than the recorded number.  Comparing the two raw lengths announced a
    one-artifact discrepancy for every normally finalized experiment.
    """
    dataset = ReportDataset(
        metadata={"experiment": {}, "template": {}},
        artifact_index=[
            {"role": "measured_values"},
            {"role": "run_results"},
            {"role": "summary_metadata"},
        ],
        summary_metadata={"artifact_count": 2},
    )

    render_experiment_metadata_section(document, dataset, tmp_path)
    text = _document_text(document)

    assert "Артефактов: 3" in text
    assert "в записи" not in text, f"a consistent archive was reported as stale:\n{text}"
